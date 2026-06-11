from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "deliverables"
ASSET_DIR = OUT_DIR / "assets"
DOCX_PATH = OUT_DIR / "西域羊都网上交易程序作品介绍.docx"
COVER_IMAGE = ASSET_DIR / "work-intro-cover.png"
FLOW_IMAGE = ASSET_DIR / "work-intro-flow.png"

FONT = "Microsoft YaHei"
FONT_PATH = Path("C:/Windows/Fonts/msyh.ttc")
BOLD_FONT_PATH = Path("C:/Windows/Fonts/msyhbd.ttc")

BLUE = RGBColor(31, 78, 121)
INK = RGBColor(32, 50, 39)
GREEN = RGBColor(40, 100, 79)
MUTED = RGBColor(92, 108, 96)
LIGHT_FILL = "F4F8EF"
TABLE_HEADER = "E8EEF5"
TABLE_BORDER = "D6DFD2"


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def style_paragraph(paragraph, before: float = 0, after: float = 6, line: float = 1.1) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_paragraph()
    style_paragraph(paragraph, before=14 if level == 1 else 8, after=6)
    run = paragraph.add_run(text)
    if level == 1:
        set_run_font(run, 16, True, BLUE)
    elif level == 2:
        set_run_font(run, 13, True, BLUE)
    else:
        set_run_font(run, 12, True, RGBColor(31, 77, 120))
    return paragraph


def add_body(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    style_paragraph(paragraph, after=6, line=1.15)
    run = paragraph.add_run(text)
    set_run_font(run, 10.5, False, INK)
    return paragraph


def add_bullet(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="List Bullet")
    style_paragraph(paragraph, after=4, line=1.15)
    run = paragraph.add_run(text)
    set_run_font(run, 10.5, False, INK)
    return paragraph


def add_step(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="List Number")
    style_paragraph(paragraph, after=4, line=1.15)
    run = paragraph.add_run(text)
    set_run_font(run, 10.5, False, INK)
    return paragraph


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: RGBColor = INK, size: float = 9.5) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    style_paragraph(paragraph, after=0, line=1.1)
    run = paragraph.add_run(text)
    set_run_font(run, size, bold, color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_margins(cell, top: int = 100, start: int = 120, bottom: int = 100, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = TABLE_BORDER) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{border_name}"))
        if border is None:
            border = OxmlElement(f"w:{border_name}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)


def set_table_widths(table, widths_in: list[float]) -> None:
    for row in table.rows:
        for cell, width in zip(row.cells, widths_in):
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            set_cell_margins(cell)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_in: list[float]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_borders(table)
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, True, BLUE, 9.5)
        set_cell_shading(table.rows[0].cells[idx], TABLE_HEADER)
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            set_cell_text(cells[idx], value)
    set_table_widths(table, widths_in)
    doc.add_paragraph()
    return table


def add_callout(doc: Document, label: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_borders(table, color="C9D8CE")
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_FILL)
    set_cell_margins(cell, top=160, bottom=160, start=180, end=180)
    paragraph = cell.paragraphs[0]
    style_paragraph(paragraph, after=3, line=1.12)
    r1 = paragraph.add_run(label)
    set_run_font(r1, 10, True, GREEN)
    r2 = paragraph.add_run(f"  {text}")
    set_run_font(r2, 10, False, INK)
    set_table_widths(table, [6.35])
    doc.add_paragraph()


def draw_flow_image() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    width, height = 1800, 780
    image = Image.new("RGB", (width, height), "#F5F7F0")
    draw = ImageDraw.Draw(image)
    font_title = ImageFont.truetype(str(BOLD_FONT_PATH), 54)
    font_block = ImageFont.truetype(str(BOLD_FONT_PATH), 34)
    font_small = ImageFont.truetype(str(FONT_PATH), 25)
    draw.text((72, 52), "系统功能结构", fill="#203227", font=font_title)
    draw.text((72, 122), "前端负责订单交互，后端负责接收、计算、归档与运营监控", fill="#5C6C60", font=font_small)
    blocks = [
        ("顾客前端商城", "选品 / 购物车 / 下单"),
        ("FastAPI 后端", "订单校验 / 金额计算 / 接口服务"),
        ("SQLite 数据库", "商品 / 订单 / 明细 / 履约数据"),
        ("汇总归纳模块", "GMV / 客单价 / 渠道 / 库存"),
        ("运营数据中心", "订单监控 / 自动刷新 / 结构分析"),
    ]
    x_positions = [72, 405, 738, 1071, 1404]
    y = 255
    block_w, block_h = 260, 210
    for i, ((title, subtitle), x) in enumerate(zip(blocks, x_positions)):
        fill = "#FFFFFF" if i % 2 == 0 else "#EAF4ED"
        outline = "#BFD6C7"
        draw.rounded_rectangle((x, y, x + block_w, y + block_h), radius=28, fill=fill, outline=outline, width=4)
        draw.text((x + 28, y + 45), title, fill="#1F4D3B", font=font_block)
        lines = subtitle.split(" / ")
        for line_idx, line in enumerate(lines):
            draw.text((x + 28, y + 112 + line_idx * 34), line, fill="#5C6C60", font=font_small)
        if i < len(blocks) - 1:
            x1 = x + block_w + 18
            x2 = x_positions[i + 1] - 18
            yy = y + block_h // 2
            draw.line((x1, yy, x2, yy), fill="#2F735A", width=6)
            draw.polygon([(x2, yy), (x2 - 20, yy - 12), (x2 - 20, yy + 12)], fill="#2F735A")
    draw.rounded_rectangle((72, 560, 1728, 690), radius=24, fill="#FFFFFF", outline="#D6DFD2", width=3)
    note = "数据闭环：前端提交订单 → 后端写入并计算 → 汇总模块归纳 → 前端与后台按版本自动刷新"
    draw.text((120, 602), note, fill="#203227", font=font_small)
    image.save(FLOW_IMAGE, quality=95)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15


def add_cover(doc: Document) -> None:
    title = doc.add_paragraph()
    style_paragraph(title, after=4)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("西域羊都网上交易程序")
    set_run_font(run, 24, True, RGBColor(18, 39, 29))

    subtitle = doc.add_paragraph()
    style_paragraph(subtitle, after=10)
    run = subtitle.add_run("作品介绍文档")
    set_run_font(run, 15, False, GREEN)

    meta = doc.add_paragraph()
    style_paragraph(meta, after=12)
    run = meta.add_run("内容范围：作品开发平台、主要软件、操作方法、作品亮点与特色")
    set_run_font(run, 10.5, False, MUTED)

    if COVER_IMAGE.exists():
        doc.add_picture(str(COVER_IMAGE), width=Inches(6.35))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_callout(
        doc,
        "作品定位",
        "本作品以西北羊产业数字化供应链为业务背景，构建了前端顾客交易、后端数据归纳与运营监控协同的网上交易程序。",
    )

    info = add_table(
        doc,
        ["项目项", "说明"],
        [
            ["作品名称", "西域羊都网上交易程序"],
            ["作品类型", "前后端分离的电商交易与运营数据监控系统"],
            ["当前访问", "前端 http://127.0.0.1:5173/；后端 http://127.0.0.1:8000/"],
            ["交付形式", "React 前端、FastAPI 后端、SQLite 数据库、Word 作品介绍文档"],
        ],
        [1.7, 4.65],
    )
    info.style = "Table Grid"
    doc.add_section(WD_SECTION.NEW_PAGE)


def build_doc() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    draw_flow_image()

    doc = Document()
    configure_document(doc)
    add_cover(doc)

    add_heading(doc, "一、作品概述")
    add_body(
        doc,
        "西域羊都网上交易程序围绕“产地直供、数字溯源、冷链履约、订单归纳”四个核心环节展开。"
        "系统既提供面向顾客的商城界面，也提供面向运营人员的数据中心，用同一套后端订单与汇总数据支撑前台交易和后台监控。",
    )
    add_bullet(doc, "前端侧重顾客视角：商品浏览、购物车、下单确认、订单状态自动更新。")
    add_bullet(doc, "后端侧重运营视角：订单接收、金额计算、数据归档、渠道结构、库存与履约监控。")
    add_bullet(doc, "汇总侧重管理视角：将订单、商品、冷链、金融和养殖户收益数据归纳为可读指标。")

    add_heading(doc, "二、作品开发平台")
    add_table(
        doc,
        ["平台要素", "采用方案", "作用"],
        [
            ["开发环境", "Windows 本地工作区", "完成前后端代码开发、运行、测试与文档生成"],
            ["前端运行平台", "Node.js + Vite 开发服务器", "启动商城界面，支持热更新与生产构建"],
            ["后端运行平台", "Python + FastAPI + Uvicorn", "提供 REST 接口、后台运营数据中心和业务计算能力"],
            ["数据平台", "SQLite 本地数据库", "保存商品、订单、订单明细、履约、金融与收益样例数据"],
            ["版本管理", "Git + GitHub", "记录功能迭代并同步远端代码仓库"],
        ],
        [1.25, 2.05, 3.05],
    )

    add_heading(doc, "三、主要软件")
    add_table(
        doc,
        ["软件/框架", "所在模块", "主要用途"],
        [
            ["React", "前端", "构建商品展示、购物车、订单确认和状态刷新交互"],
            ["Vite", "前端", "提供开发服务器与生产构建能力"],
            ["FastAPI", "后端", "定义商品、订单、汇总、变更状态和溯源接口"],
            ["Uvicorn", "后端", "运行 ASGI Web 服务"],
            ["SQLite", "后端数据", "存储样例交易和经营数据"],
            ["Python unittest / Node test", "测试", "验证后端汇总、后台页面约束和前端计算逻辑"],
        ],
        [1.5, 1.35, 3.5],
    )

    add_heading(doc, "四、系统结构与数据流")
    if FLOW_IMAGE.exists():
        doc.add_picture(str(FLOW_IMAGE), width=Inches(6.35))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_body(
        doc,
        "系统采用前后端分离结构。前端只负责顾客可见的购买体验，后端负责接收订单、校验商品、计算金额、写入数据库，并将经营数据归纳为运营指标。"
        "前端和后台都通过订单变更版本进行轻量轮询，避免频繁拉取全量数据。",
    )

    add_heading(doc, "五、操作方法")
    add_heading(doc, "1. 启动后端服务", 2)
    add_step(doc, "进入 backend 目录，安装依赖：python -m pip install -r requirements.txt。")
    add_step(doc, "运行后端：uvicorn app.main:app --reload --port 8000。")
    add_step(doc, "打开 http://127.0.0.1:8000/ 查看运营数据中心；打开 /docs 可查看接口文档。")
    add_heading(doc, "2. 启动前端商城", 2)
    add_step(doc, "进入 frontend 目录，安装依赖：npm install。")
    add_step(doc, "运行前端：npm run dev。")
    add_step(doc, "打开 http://127.0.0.1:5173/，从顾客视角完成选品、加入购物车和提交订单。")
    add_heading(doc, "3. 体验完整交易流程", 2)
    add_step(doc, "在前端商城选择商品并调整数量，填写采购方名称和采购场景。")
    add_step(doc, "点击提交订单，后端接收订单并按商品单价计算订单总额。")
    add_step(doc, "返回后端运营数据中心，查看订单监控、渠道结构、库存、履约与数据版本是否自动刷新。")

    add_heading(doc, "六、作品亮点与特色")
    add_table(
        doc,
        ["亮点", "体现方式", "价值"],
        [
            ["前后端职责清晰", "前端负责交易交互，后端负责数据计算和运营查看", "符合真实业务分工，避免后台承担顾客下单职责"],
            ["真实订单闭环", "POST /api/orders 写入订单，summary 模块重新归纳交易指标", "证明前端改变可以牵引后端数据变化"],
            ["自动刷新机制", "前端和后台每 5 秒检查 /api/change-state", "减少人工刷新，模拟实时交易场景"],
            ["后台运营数据中心", "展示订单监控、履约、商品库存、渠道与客户结构", "让后端成为管理者的数据查看入口"],
            ["数据清洗意识", "对历史乱码标签进行归一化展示", "提升展示可信度和真实系统健壮性"],
            ["产业场景贴合", "溯源、冷链、库存、渠道和客户结构共同呈现", "体现西北羊产业数字化供应链特色"],
        ],
        [1.45, 2.5, 2.4],
    )

    add_heading(doc, "七、测试与实现程度")
    add_body(
        doc,
        "项目已设置前后端测试：后端覆盖综合汇总、订单变更状态、后台控制台约束和标签清洗；前端覆盖购物车金额、商品推荐、下单条件和自动刷新判断。"
        "生产构建用于验证前端可发布性，接口验证用于确认后端能接收订单并更新汇总。",
    )
    add_table(
        doc,
        ["验证项", "当前证据", "说明"],
        [
            ["后端单元测试", "unittest discover -s backend/tests", "覆盖数据汇总、后台页面和变更状态"],
            ["前端单元测试", "npm test", "覆盖购物车、提交条件与刷新判断"],
            ["前端生产构建", "npm run build", "验证 Vite 项目可构建"],
            ["接口联动", "前端提交订单后，后端订单数和交易额变化", "证明前后端数据流可行"],
        ],
        [1.55, 2.3, 2.5],
    )

    add_heading(doc, "八、作品总结")
    add_body(
        doc,
        "本作品不是单纯的静态展示页面，而是一个可运行、可下单、可归纳、可监控的交易程序。"
        "它将西域羊都的产业背景转化为软件系统中的商品、订单、履约和经营数据，并通过前端商城与后端运营数据中心分别服务顾客和管理人员。"
        "作品亮点在于业务逻辑完整、交互路径清晰、后端具备真实计算处理能力，且能够根据前端订单变化及时刷新数据。",
    )

    footer = doc.sections[-1].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("西域羊都网上交易程序 · 作品介绍")
    set_run_font(run, 8.5, False, MUTED)

    doc.core_properties.title = "西域羊都网上交易程序作品介绍"
    doc.core_properties.subject = "作品开发平台、主要软件、操作方法及亮点特色"
    doc.core_properties.keywords = "React, FastAPI, SQLite, 网上交易程序, 作品介绍"
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
